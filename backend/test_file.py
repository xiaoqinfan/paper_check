import os
from docx import Document

# 构建绝对路径
base_dir = os.path.abspath(os.path.dirname(__file__))
uploads_dir = os.path.join(base_dir, "uploads")

print(f"Base directory: {base_dir}")
print(f"Uploads directory: {uploads_dir}")

# 检查uploads目录是否存在
if not os.path.exists(uploads_dir):
    print(f"Uploads directory does not exist: {uploads_dir}")
else:
    print("Uploads directory exists")
    
    # 列出uploads目录中的所有文件
    import glob
    files = glob.glob(os.path.join(uploads_dir, "*"))
    print(f"Files in uploads directory: {files}")
    
    # 如果有文件，尝试打开第一个文件
    if files:
        file_path = files[0]
        print(f"Using first file: {file_path}")
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            print(f"File does not exist: {file_path}")
        else:
            print("File exists")
            
            # 尝试打开并解析论文文件
            try:
                print(f"Attempting to open document: {file_path}")
                doc = Document(file_path)
                print("Successfully opened document")
                
                # 打印文档的段落数
                print(f"Number of paragraphs: {len(doc.paragraphs)}")
                
                # 打印前几个段落的内容
                for i, para in enumerate(doc.paragraphs[:5]):
                    print(f"Paragraph {i}: {para.text}")
            except Exception as e:
                print(f"Error opening document: {str(e)}")
                import traceback
                traceback.print_exc()
    else:
        print("No files in uploads directory")
