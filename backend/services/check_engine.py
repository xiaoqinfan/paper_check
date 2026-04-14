# check_engine.py（论文格式检测核心）
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
import re

# 规则示例（可从数据库format_rule_item读取）
def check_paper(docx_path, rules):
    doc = Document(docx_path)
    errors = []
    score = 100
    error_count = 0

    # 1. 检测全文字体
    for para in doc.paragraphs:
        for run in para.runs:
            # 规则：正文 宋体 12pt
            rule = next((r for r in rules if r["target_part"]=="正文" and r["rule_key"]=="font_name"), None)
            if rule and run.font.name != rule["rule_value"]:
                errors.append({
                    "position": f"段落{para.text[:20]}...",
                    "error": "字体错误",
                    "expect": rule["rule_value"],
                    "actual": run.font.name
                })
                error_count += 1
                score -= 2

            # 字号
            rule_size = next((r for r in rules if r["target_part"]=="正文" and r["rule_key"]=="font_size"), None)
            if rule_size and run.font.size.pt != float(rule_size["rule_value"]):
                errors.append({
                    "position": f"段落{para.text[:20]}...",
                    "error": "字号错误",
                    "expect": rule_size["rule_value"],
                    "actual": run.font.size.pt
                })
                error_count +=1
                score -=2

    # 2. 参考文献 GB/T 7714 检测
    ref_pattern = re.compile(r'^\[\d+\].*')
    for i, para in enumerate(doc.paragraphs):
        if ref_pattern.match(para.text):
            # 期刊：[1]作者.题名[J].刊名,年,卷(期):起止页码.
            if not re.search(r'\[J\]', para.text):
                errors.append({
                    "position": f"参考文献第{i+1}条",
                    "error": "参考文献格式错误（缺少[J]）",
                    "expect": "GB/T 7714 期刊格式",
                    "actual": para.text[:50]
                })
                error_count +=1
                score -=5

    return {
        "score": max(score, 0),
        "error_count": error_count,
        "errors": errors
    }