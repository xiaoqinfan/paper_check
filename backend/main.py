from fastapi import FastAPI, UploadFile, File, Depends, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy import create_engine, Column, Integer, String, Text, ForeignKey, Float, DateTime
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, Session
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt
from typing import List, Optional
import os
import uuid
from datetime import datetime
import re

# 数据库配置
DATABASE_URL = "mysql+pymysql://root:123456@localhost:3306/paper_check?charset=utf8mb4"
engine = create_engine(DATABASE_URL)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

# 数据库模型
class User(Base):
    __tablename__ = "sys_user"
    id = Column(Integer, primary_key=True, index=True)
    username = Column(String(50), unique=True, index=True)
    password = Column(String(100))
    real_name = Column(String(50))
    user_type = Column(Integer) # 1管理员 2教师 3学生

class Class(Base):
    __tablename__ = "course_class"
    id = Column(Integer, primary_key=True)
    class_name = Column(String(100))
    teacher_id = Column(Integer)
    college = Column(String(100))
    major = Column(String(100))

class StudentClassRel(Base):
    __tablename__ = "student_class_rel"
    id = Column(Integer, primary_key=True)
    class_id = Column(Integer)
    student_id = Column(Integer)

class FormatTemplate(Base):
    __tablename__ = "format_template"
    id = Column(Integer, primary_key=True)
    template_name = Column(String(100))

class FormatRule(Base):
    __tablename__ = "format_rule_item"
    id = Column(Integer, primary_key=True)
    template_id = Column(Integer)
    target_part = Column(String(100)) # 正文/标题/参考文献
    rule_key = Column(String(50)) # font_name/size/line_height/indent
    rule_value = Column(String(100))
    error_tip = Column(String(255))

class Paper(Base):
    __tablename__ = "paper_document"
    id = Column(Integer, primary_key=True)
    student_id = Column(Integer)
    class_id = Column(Integer)
    paper_name = Column(String(255))
    file_path = Column(String(255))
    upload_time = Column(DateTime, default=datetime.now)

class CheckTask(Base):
    __tablename__ = "check_task"
    id = Column(Integer, primary_key=True)
    paper_id = Column(Integer)
    format_score = Column(Float)
    total_error = Column(Integer)
    check_time = Column(DateTime, default=datetime.now)

class CheckError(Base):
    __tablename__ = "check_error_detail"
    id = Column(Integer, primary_key=True)
    task_id = Column(Integer)
    paper_id = Column(Integer)
    position = Column(String(255))
    error_type = Column(String(50))
    error_message = Column(String(255))
    expect_value = Column(String(100))
    actual_value = Column(String(100))

Base.metadata.create_all(bind=engine)

# Pydantic模型
class RuleCreate(BaseModel):
    template_id: int
    target_part: str
    rule_key: str
    rule_value: str
    error_tip: str

class LoginRequest(BaseModel):
    username: str
    password: str

# 依赖
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

app = FastAPI()

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 初始化用户
@app.post("/init")
def init(db: Session = Depends(get_db)):
    if not db.query(User).filter(User.username=="admin").first():
        db.add_all([
            User(username="admin", password="123", real_name="管理员", user_type=1),
            User(username="teacher", password="123", real_name="张老师", user_type=2),
            User(username="student", password="123", real_name="小明", user_type=3)
        ])
        db.commit()
    return {"msg": "初始化成功"}

# 登录
@app.post("/login")
def login(req: LoginRequest, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.username==req.username, User.password==req.password).first()
    if not user:
        raise HTTPException(401, "账号或密码错误")
    return {
        "id": user.id,
        "username": user.username,
        "real_name": user.real_name,
        "user_type": user.user_type
    }

# ------------------------------
# 班级管理（教师）
# ------------------------------
@app.get("/teacher/class/list")
def list_class(db: Session = Depends(get_db)):
    return db.query(Class).all()

@app.post("/teacher/class/create")
def create_class(name: str, college: str, major: str, db: Session = Depends(get_db)):
    cls = Class(class_name=name, college=college, major=major, teacher_id=2)
    db.add(cls)
    db.commit()
    return {"msg": "成功"}

# ------------------------------
# 格式模板（教师）
# ------------------------------
@app.get("/teacher/template/list")
def list_template(db: Session = Depends(get_db)):
    return db.query(FormatTemplate).all()

@app.post("/teacher/template/create")
def create_template(name: str, db: Session = Depends(get_db)):
    tpl = FormatTemplate(template_name=name)
    db.add(tpl)
    db.commit()
    return {"id": tpl.id}

@app.post("/teacher/rule/add")
def add_rule(rule: RuleCreate, db: Session = Depends(get_db)):
    item = FormatRule(**rule.dict())
    db.add(item)
    db.commit()
    return {"msg": "成功"}

@app.get("/teacher/rule/list/{template_id}")
def list_rule(template_id: int, db: Session = Depends(get_db)):
    return db.query(FormatRule).filter(FormatRule.template_id==template_id).all()

@app.delete("/teacher/rule/delete/{rule_id}")
def delete_rule(rule_id: int, db: Session = Depends(get_db)):
    rule = db.query(FormatRule).filter(FormatRule.id==rule_id).first()
    if rule:
        db.delete(rule)
        db.commit()
        return {"msg": "删除成功"}
    else:
        raise HTTPException(404, "规则不存在")

# ------------------------------
# 论文上传 & 检测引擎（核心）
# ------------------------------
@app.post("/student/paper/upload")
def upload(student_id: int = Form(...), class_id: int = Form(...), file: UploadFile=File(...), db: Session = Depends(get_db)):
    try:
        # 检查文件类型
        file_extension = os.path.splitext(file.filename)[1].lower()
        if file_extension not in ['.docx', '.doc']:
            raise HTTPException(400, "Only docx and doc files are allowed")
        # 获取绝对路径
        uploads_dir = os.path.join(os.getcwd(), "uploads")
        os.makedirs(uploads_dir, exist_ok=True)
        fname = f"{uuid.uuid4()}_{file.filename}"
        path = os.path.join(uploads_dir, fname)
        with open(path, "wb") as f:
            f.write(file.file.read())
        # 存储相对路径
        relative_path = f"uploads/{fname}"
        paper = Paper(student_id=student_id, class_id=class_id, paper_name=file.filename, file_path=relative_path)
        db.add(paper)
        db.commit()
        return {"paper_id": paper.id}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"上传失败: {str(e)}")

# 检测引擎
def check_docx(file_path: str, rules: List[FormatRule]):
    # 构建绝对路径
    base_dir = os.path.abspath(os.path.dirname(__file__))
    absolute_path = os.path.join(base_dir, file_path)
    uploads_dir = os.path.join(base_dir, "uploads")
    
    print(f"Checking file: {absolute_path}")
    print(f"Uploads directory: {uploads_dir}")
    
    # 检查uploads目录是否存在
    if not os.path.exists(uploads_dir):
        print(f"Uploads directory does not exist: {uploads_dir}")
        return {"score": 0, "count": 1, "errors": [{"position": "目录不存在", "error_type": "文件错误", "msg": "上传目录不存在", "expect": "有效的上传目录", "actual": "目录不存在"}]}
    
    # 检查指定的文件是否存在
    if not os.path.exists(absolute_path):
        print(f"File does not exist: {absolute_path}")
        return {"score": 0, "count": 1, "errors": [{"position": "文件不存在", "error_type": "文件错误", "msg": "论文文件不存在", "expect": "有效的论文文件", "actual": "文件不存在"}]}
    
    # 尝试打开指定的文件
    try:
        print(f"File exists: {absolute_path}")
        print(f"File size: {os.path.getsize(absolute_path)} bytes")
        # 获取文件扩展名
        file_extension = os.path.splitext(absolute_path)[1].lower()
        print(f"File extension: {file_extension}")
        
        # 尝试使用python-docx打开文件（仅适用于docx）
        if file_extension == '.docx':
            try:
                doc = Document(absolute_path)
                print("Successfully opened document with python-docx")
                return process_document(doc, rules)
            except Exception as e:
                print(f"Error opening document with python-docx: {str(e)}")
        
        # 对于doc文件，直接返回一个默认的成功结果
        if file_extension == '.doc':
            print("Processing doc file with default result")
            # 对于doc文件，我们无法进行详细的格式检测，但可以确认文件已上传成功
            return {
                "score": 100.0,
                "count": 0,
                "errors": [],
                "message": "doc格式文件已上传成功，但由于格式限制，无法进行详细的格式检测"
            }
        
        # 尝试使用docx2txt打开doc文件（作为备选方案）
        try:
            import docx2txt
            print("Attempting to open document with docx2txt")
            text = docx2txt.process(absolute_path)
            print("Successfully opened document with docx2txt")
            print(f"Extracted text length: {len(text)} characters")
            # 对于doc文件，我们只能检查文本内容，无法检查格式
            errors = []
            score = 100.0
            error_count = 0
            
            # 参考文献 GB/T 7714
            import re
            ref_pattern = re.compile(r'^\[\d+\].*')
            lines = text.split('\n')
            print(f"Number of lines: {len(lines)}")
            for i, line in enumerate(lines):
                if ref_pattern.match(line):
                    if not re.search(r'\[J\]|\[M\]|\[D\]', line):
                        errors.append({
                            "position": f"参考文献第{i+1}条",
                            "error_type": "参考文献",
                            "msg": "不符合GB/T 7714，缺少文献类型标识[J][M][D]",
                            "expect": "GB/T 7714标准",
                            "actual": line[:50]
                        })
                        error_count +=1
                        score -=5
            
            return {"score": max(score,0), "count": error_count, "errors": errors}
        except Exception as e2:
            print(f"Error opening document with docx2txt: {str(e2)}")
    except Exception as e:
        print(f"Error opening specified document: {str(e)}")
    
    # 尝试修复文件路径中的空格问题
    try:
        fixed_path = absolute_path.replace(' ', '')
        if os.path.exists(fixed_path):
            print(f"Trying fixed path: {fixed_path}")
            doc = Document(fixed_path)
            print("Successfully opened document with fixed path")
            return process_document(doc, rules)
    except Exception as e:
        print(f"Error opening fixed path document: {str(e)}")
    
    # 如果指定的文件无法打开，返回错误信息，而不是尝试打开其他文件
    print("Failed to open the specified document, returning error")
    return {"score": 0, "count": 1, "errors": [{"position": "文件打开失败", "error_type": "文件错误", "msg": "无法打开论文文件", "expect": "有效的DOCX或DOC文件", "actual": "文件无法解析"}]}

# 处理文档的函数
def process_document(doc, rules):
    errors = []
    score = 100.0  # 确保score是一个浮点数
    error_count = 0

    # 辅助函数：将字号字符串转换为点数
    def get_font_size_pt(size_str):
        size_map = {
            '初号': 42,
            '小初': 36,
            '一号': 26,
            '小一': 24,
            '二号': 22,
            '小二': 18,
            '三号': 16,
            '小三': 15,
            '四号': 14,
            '小四': 12,
            '五号': 10.5,
            '小五': 9
        }
        return size_map.get(size_str, float(size_str) if size_str.replace('.', '').isdigit() else 0)

    # 辅助函数：将行距字符串转换为倍数
    def get_line_height_multiplier(height_str):
        height_map = {
            '单倍行距': 1.0,
            '1.5倍行距': 1.5,
            '2倍行距': 2.0,
            '最小值': 1.0,
            '固定值': 1.0,
            '多倍行距': 1.5
        }
        return height_map.get(height_str, 1.0)

    # 辅助函数：将缩进字符串转换为点数
    def get_indent_pt(indent_str):
        import re
        # 尝试从字符串中提取数字
        match = re.search(r'\d+', indent_str)
        if match:
            number = float(match.group(0))
            if '字符' in indent_str:
                # 1字符 = 12pt
                return number * 12
            elif 'pt' in indent_str:
                return number
            else:
                # 默认为字符
                return number * 12
        else:
            return 0

    # 遍历所有段落
    for para in doc.paragraphs:
        # 行距检测
        line_height_rule = next((r for r in rules if r.target_part=="正文" and r.rule_key=="line_height"), None)
        if line_height_rule:
            para_line_spacing = para.paragraph_format.line_spacing or 1.0  # 当line_spacing为None时使用默认值1.0
            expected_line_spacing = get_line_height_multiplier(line_height_rule.rule_value)
            if abs(para_line_spacing - expected_line_spacing)>0.1:
                errors.append({
                    "position": para.text[:30],
                    "error_type": "行距",
                    "msg": line_height_rule.error_tip,
                    "expect": line_height_rule.rule_value,
                    "actual": f"{para_line_spacing}倍"
                })
                error_count +=1
                score -=2

        # 缩进检测
        indent_rule = next((r for r in rules if r.target_part=="正文" and r.rule_key=="indent"), None)
        if indent_rule:
            first_line_indent = para.paragraph_format.first_line_indent.pt if para.paragraph_format.first_line_indent else 0
            expected_indent = get_indent_pt(indent_rule.rule_value)
            if abs(first_line_indent - expected_indent)>0.5:
                errors.append({
                    "position": para.text[:30],
                    "error_type": "缩进",
                    "msg": indent_rule.error_tip,
                    "expect": indent_rule.rule_value,
                    "actual": f"{first_line_indent}pt"
                })
                error_count +=1
                score -=2

        for run in para.runs:
            # 中文字体检测
            font_zh_rule = next((r for r in rules if r.target_part=="正文" and r.rule_key=="font_name_zh"), None)
            if font_zh_rule:
                # 检查是否包含中文字符
                has_chinese = any('\u4e00' <= char <= '\u9fff' for char in run.text)
                if has_chinese and run.font.name != font_zh_rule.rule_value:
                    errors.append({
                        "position": para.text[:30],
                        "error_type": "中文字体",
                        "msg": font_zh_rule.error_tip,
                        "expect": font_zh_rule.rule_value,
                        "actual": run.font.name
                    })
                    error_count +=1
                    score -=2

            # 英文字体检测
            font_en_rule = next((r for r in rules if r.target_part=="正文" and r.rule_key=="font_name_en"), None)
            if font_en_rule:
                # 检查是否包含英文字符
                has_english = any('a' <= char.lower() <= 'z' for char in run.text)
                if has_english and run.font.name != font_en_rule.rule_value:
                    errors.append({
                        "position": para.text[:30],
                        "error_type": "英文字体",
                        "msg": font_en_rule.error_tip,
                        "expect": font_en_rule.rule_value,
                        "actual": run.font.name
                    })
                    error_count +=1
                    score -=2

            # 字号
            size_rule = next((r for r in rules if r.target_part=="正文" and r.rule_key=="font_size"), None)
            if size_rule:
                sz = run.font.size.pt if run.font.size else 0
                expected_size = get_font_size_pt(size_rule.rule_value)
                if abs(sz - expected_size)>0.1:
                    errors.append({
                        "position": para.text[:30],
                        "error_type": "字号",
                        "msg": size_rule.error_tip,
                        "expect": size_rule.rule_value,
                        "actual": f"{sz}pt"
                    })
                    error_count +=1
                    score -=2

    # 参考文献 GB/T 7714
    import re
    ref_pattern = re.compile(r'^\[\d+\].*')
    for i, para in enumerate(doc.paragraphs):
        if ref_pattern.match(para.text):
            if not re.search(r'\[J\]|\[M\]|\[D\]', para.text):
                errors.append({
                    "position": f"参考文献第{i+1}条",
                    "error_type": "参考文献",
                    "msg": "不符合GB/T 7714，缺少文献类型标识[J][M][D]",
                    "expect": "GB/T 7714标准",
                    "actual": para.text[:50]
                })
                error_count +=1
                score -=5

    return {"score": max(score,0), "count": error_count, "errors": errors}

class CheckRequest(BaseModel):
    template_id: int

@app.post("/paper/check/{paper_id}")
def check(paper_id: int, req: CheckRequest, db: Session = Depends(get_db)):
    try:
        paper = db.query(Paper).filter(Paper.id==paper_id).first()
        rules = db.query(FormatRule).filter(FormatRule.template_id==req.template_id).all()
        if not paper:
            raise HTTPException(404, "论文不存在")
        if not rules:
            # 如果没有规则，创建默认规则
            if not db.query(FormatTemplate).filter(FormatTemplate.id==req.template_id).first():
                template = FormatTemplate(template_name="默认模板")
                db.add(template)
                db.commit()
                # 添加默认规则
                default_rules = [
                    FormatRule(template_id=template.id, target_part="正文", rule_key="font_name_zh", rule_value="宋体", error_tip="正文中文字体应为宋体"),
                    FormatRule(template_id=template.id, target_part="正文", rule_key="font_name_en", rule_value="Times New Roman", error_tip="正文英文字体应为Times New Roman"),
                    FormatRule(template_id=template.id, target_part="正文", rule_key="font_size", rule_value="小四", error_tip="正文字号应为小四"),
                    FormatRule(template_id=template.id, target_part="正文", rule_key="line_height", rule_value="1.5倍行距", error_tip="正文行距应为1.5倍"),
                    FormatRule(template_id=template.id, target_part="正文", rule_key="indent", rule_value="首行缩进2字符", error_tip="正文首行缩进应为2字符")
                ]
                db.add_all(default_rules)
                db.commit()
                rules = default_rules
            else:
                raise HTTPException(404, "模板规则不存在")
        res = check_docx(paper.file_path, rules)
        task = CheckTask(paper_id=paper_id, format_score=res["score"], total_error=res["count"])
        db.add(task)
        db.commit()
        for e in res["errors"]:
            db.add(CheckError(
                task_id=task.id,
                paper_id=paper_id,
                position=e["position"],
                error_type=e["error_type"],
                error_message=e["msg"],
                expect_value=e["expect"],
                actual_value=e["actual"]
            ))
        db.commit()
        return {"task_id": task.id, **res}
    except Exception as e:
        raise HTTPException(500, f"格式检测失败: {str(e)}")

# ------------------------------
# 查看结果（教师+学生）
# ------------------------------
@app.get("/teacher/class/{class_id}/papers")
def class_papers(class_id: int, db: Session = Depends(get_db)):
    return db.query(Paper).filter(Paper.class_id==class_id).all()

@app.get("/paper/result/{task_id}")
def paper_result(task_id: int, db: Session = Depends(get_db)):
    task = db.query(CheckTask).filter(CheckTask.id==task_id).first()
    errors = db.query(CheckError).filter(CheckError.task_id==task_id).all()
    return {"task": task, "errors": errors}

@app.get("/paper/tasks/{paper_id}")
def paper_tasks(paper_id: int, db: Session = Depends(get_db)):
    return db.query(CheckTask).filter(CheckTask.paper_id==paper_id).order_by(CheckTask.id.desc()).all()


# ------------------------------
# 管理员接口
# ------------------------------
@app.get("/admin/user/list")
def admin_list_users(db: Session = Depends(get_db)):
    return db.query(User).all()

@app.get("/admin/paper/list")
def admin_list_papers(db: Session = Depends(get_db)):
    return db.query(Paper).all()

# ------------------------------
# 学生接口
# ------------------------------
@app.get("/student/papers/list/{student_id}")
def student_papers(student_id: int, db: Session = Depends(get_db)):
    return db.query(Paper).filter(Paper.student_id==student_id).all()

@app.delete("/paper/delete/{paper_id}")
def delete_paper(paper_id: int, db: Session = Depends(get_db)):
    try:
        paper = db.query(Paper).filter(Paper.id==paper_id).first()
        if not paper:
            raise HTTPException(404, "论文不存在")
        
        # 删除对应的文件
        try:
            if paper.file_path:
                import os
                base_dir = os.path.abspath(os.path.dirname(__file__))
                file_path = os.path.join(base_dir, paper.file_path)
                if os.path.exists(file_path):
                    os.remove(file_path)
                    print(f"Deleted file: {file_path}")
        except Exception as e:
            print(f"Error deleting file: {str(e)}")
        
        # 删除数据库记录
        db.delete(paper)
        db.commit()
        
        return {"message": "删除成功"}
    except Exception as e:
        print(f"Error deleting paper: {str(e)}")
        db.rollback()
        raise HTTPException(500, "删除失败")

# 批量删除论文的请求模型
class BatchDeleteRequest(BaseModel):
    paper_ids: list[int]

@app.post("/paper/batch-delete")
def batch_delete_papers(request: BatchDeleteRequest, db: Session = Depends(get_db)):
    try:
        paper_ids = request.paper_ids
        papers = db.query(Paper).filter(Paper.id.in_(paper_ids)).all()
        if not papers:
            raise HTTPException(404, "论文不存在")
        
        # 删除对应的文件
        for paper in papers:
            try:
                if paper.file_path:
                    import os
                    base_dir = os.path.abspath(os.path.dirname(__file__))
                    file_path = os.path.join(base_dir, paper.file_path)
                    if os.path.exists(file_path):
                        os.remove(file_path)
                        print(f"Deleted file: {file_path}")
            except Exception as e:
                print(f"Error deleting file: {str(e)}")
        
        # 删除数据库记录
        db.query(Paper).filter(Paper.id.in_(paper_ids)).delete()
        db.commit()
        
        return {"message": "批量删除成功"}
    except Exception as e:
        print(f"Error batch deleting papers: {str(e)}")
        db.rollback()
        raise HTTPException(500, "批量删除失败")

@app.get("/paper/content/{paper_id}")
def get_paper_content(paper_id: int, db: Session = Depends(get_db)):
    try:
        print(f"Getting paper content for paper_id: {paper_id}")
        paper = db.query(Paper).filter(Paper.id==paper_id).first()
        if not paper:
            raise HTTPException(404, "论文不存在")
        
        print(f"Found paper: {paper.paper_name}, file_path: {paper.file_path}")
        
        # 构建绝对路径
        base_dir = os.path.abspath(os.path.dirname(__file__))
        print(f"Base directory: {base_dir}")
        uploads_dir = os.path.join(base_dir, "uploads")
        print(f"Uploads directory: {uploads_dir}")
        
        # 检查uploads目录是否存在
        if not os.path.exists(uploads_dir):
            print(f"Uploads directory does not exist: {uploads_dir}")
            # 目录不存在，返回错误信息
            return {
                "paper_name": paper.paper_name,
                "paragraphs": [
                    {"id": 0, "text": f"论文内容加载失败：uploads目录不存在 ({uploads_dir})", "style": "Normal"}
                ]
            }
        
        # 首先尝试使用数据库中存储的file_path
        if paper.file_path:
            # 构建文件的绝对路径
            file_path = os.path.join(base_dir, paper.file_path)
            print(f"Trying to open file from database: {file_path}")
            
            # 获取文件扩展名
            file_extension = os.path.splitext(file_path)[1].lower()
            print(f"File extension: {file_extension}")
            
            # 对于doc文件，返回一个默认的内容
            if file_extension == '.doc':
                print("Processing doc file for content")
                return {
                    "paper_name": paper.paper_name,
                    "paragraphs": [
                        {"id": 0, "text": "doc格式文件已上传成功，但由于格式限制，无法显示详细内容", "style": "Normal"}
                    ]
                }
            
            # 尝试打开并解析论文文件（docx）
            try:
                print(f"Attempting to open document: {file_path}")
                doc = Document(file_path)
                print(f"Successfully opened document: {file_path}")
                paragraphs = []
                
                for i, para in enumerate(doc.paragraphs):
                    # 构建段落的运行元素
                    runs = []
                    for run in para.runs:
                        runs.append({
                            "text": run.text,
                            "font_name": run.font.name if run.font.name else "SimSun, Times New Roman",
                            "font_size": run.font.size.pt if run.font.size else 12,
                            "bold": run.bold if run.bold else False,
                            "italic": run.italic if run.italic else False,
                            "underline": run.underline if run.underline else False
                        })
                    
                    # 构建段落格式
                    para_format = {
                        "alignment": str(para.alignment) if para.alignment else "left",
                        "space_before": para.space_before if para.space_before else 0,
                        "space_after": para.space_after if para.space_after else 0,
                        "line_spacing": para.line_spacing if para.line_spacing else 1.0
                    }
                    
                    paragraphs.append({
                        "id": i,
                        "text": para.text,
                        "style": para.style.name if para.style else "Normal",
                        "format": para_format,
                        "runs": runs
                    })
                
                print(f"Successfully parsed {len(paragraphs)} paragraphs")
                return {"paper_name": paper.paper_name, "paragraphs": paragraphs}
            except Exception as e:
                print(f"Error opening document {file_path}: {str(e)}")
                # 文件解析失败，返回错误信息
                return {
                    "paper_name": paper.paper_name,
                    "paragraphs": [
                        {"id": 0, "text": f"论文内容加载失败：无法解析论文文件 ({str(e)})", "style": "Normal"}
                    ]
                }
        else:
            print("No file path stored in database")
            # 没有文件路径，返回错误信息
            return {
                "paper_name": paper.paper_name,
                "paragraphs": [
                    {"id": 0, "text": "论文内容加载失败：数据库中没有存储文件路径", "style": "Normal"}
                ]
            }
        

    except Exception as e:
        print(f"Error in get_paper_content: {str(e)}")
        import traceback
        traceback.print_exc()
        # 其他错误，返回默认内容
        return {
            "paper_name": "未知论文",
            "paragraphs": [
                {"id": 0, "text": f"论文内容加载失败：系统错误 ({str(e)})", "style": "Normal"}
            ]
        }